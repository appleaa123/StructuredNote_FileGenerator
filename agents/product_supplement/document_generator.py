"""
Document generator for PDS (Prospectus Supplement) agent outputs.

TODO: Full implementation pending ISM agent completion.
"""

import os
from datetime import datetime
from typing import Optional, Dict, Any
try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional dependency in CI
    Document = None  # type: ignore

from .models import PDSOutput, PDSInput
from .config import PDSConfig
from core.config import global_config


class PDSDocumentGenerator:
    """
    Generator for creating formatted PDS documents from structured output.
    
    TODO: Implement full document generator following ISM pattern.
    """
    
    def __init__(self, config: Optional[PDSConfig] = None):
        """
        Initialize the document generator.
        
        Args:
            config: PDS configuration for formatting preferences
        """
        self.config = config or PDSConfig.get_default_config()
        self.output_dir = global_config.get_output_path("pds")
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def save_outputs(
        self,
        pds_output: PDSOutput,
        input_data: PDSInput,
        filename_stem: Optional[str] = None,
        save_docx: bool = True,
    ) -> Dict[str, str]:
        """
        Save PDS outputs to JSON, TXT, and optionally DOCX.

        Returns dict with keys: json_path, txt_path, (optional) docx_path
        """
        import json

        if not filename_stem:
            safe_series = "".join(c for c in input_data.note_series if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename_stem = f"PDS_{safe_series}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        json_path = os.path.join(self.output_dir, f"{filename_stem}.json")
        txt_path = os.path.join(self.output_dir, f"{filename_stem}.txt")

        # JSON
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(pds_output.model_dump(), f, ensure_ascii=False, indent=2)

        # TXT
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(pds_output.document_title + "\n\n")
            f.write(pds_output.supplement_cover + "\n\n")
            f.write("Base Prospectus: " + pds_output.base_prospectus_reference + "\n\n")
            f.write("Purpose: " + pds_output.supplement_purpose + "\n\n")
            f.write("Specific Terms:\n" + pds_output.specific_terms + "\n\n")
            f.write("Underlying Description:\n" + pds_output.underlying_description + "\n\n")
            f.write("Calculation Methodology:\n" + pds_output.calculation_methodology + "\n\n")
            f.write("Payment Schedule:\n" + pds_output.payment_schedule + "\n\n")
            f.write("Additional Risks:\n" + pds_output.additional_risks + "\n\n")
            f.write("Pricing Details:\n" + pds_output.pricing_details + "\n\n")
            f.write("Market Information:\n" + pds_output.market_information + "\n\n")
            f.write("Tax Implications:\n" + pds_output.tax_implications + "\n\n")
            if pds_output.additional_sections:
                f.write("Additional Sections:\n")
                for k, v in pds_output.additional_sections.items():
                    f.write(f"[{k}]\n{v}\n\n")

        paths = {"json_path": json_path, "txt_path": txt_path}

        if save_docx:
            paths["docx_path"] = self.create_docx_document(pds_output, input_data, f"{filename_stem}.docx")

        return paths

    def create_docx_document(
        self, 
        pds_output: PDSOutput, 
        input_data: PDSInput,
        filename: Optional[str] = None
    ) -> str:
        """
        Create a formatted DOCX document from PDS output.
        
        TODO: Implement comprehensive document creation following ISM pattern.
        
        Args:
            pds_output: Structured output from PDS agent
            input_data: Original input data for context
            filename: Custom filename (optional)
            
        Returns:
            Path to the created document
        """
        if Document is None:
            raise RuntimeError("python-docx is required to create DOCX documents")
        doc = Document()

        # Title
        doc.add_heading(pds_output.document_title, level=0)

        # Cover / Offering Details
        doc.add_heading("Offering Details", level=1)
        doc.add_paragraph(pds_output.supplement_cover)

        # Base Prospectus Reference
        doc.add_heading("Base Prospectus", level=1)
        doc.add_paragraph(pds_output.base_prospectus_reference)

        # Purpose
        doc.add_heading("Purpose", level=1)
        doc.add_paragraph(pds_output.supplement_purpose)

        # Specific Terms
        doc.add_heading("Specific Terms", level=1)
        doc.add_paragraph(pds_output.specific_terms)

        # Underlying Description
        doc.add_heading("Underlying Description", level=1)
        doc.add_paragraph(pds_output.underlying_description)

        # Calculation Methodology
        doc.add_heading("Calculation Methodology", level=1)
        doc.add_paragraph(pds_output.calculation_methodology)

        # Payment Schedule
        doc.add_heading("Payment Schedule", level=1)
        doc.add_paragraph(pds_output.payment_schedule)

        # Risks
        doc.add_heading("Risk Factors", level=1)
        doc.add_paragraph(pds_output.additional_risks)

        # Pricing
        doc.add_heading("Pricing Details", level=1)
        doc.add_paragraph(pds_output.pricing_details)

        # Market Information
        doc.add_heading("Market Information", level=1)
        doc.add_paragraph(pds_output.market_information)

        # Tax
        doc.add_heading("Tax Implications", level=1)
        doc.add_paragraph(pds_output.tax_implications)

        # Additional Sections
        if pds_output.additional_sections:
            doc.add_heading("Additional Sections", level=1)
            for k, v in pds_output.additional_sections.items():
                doc.add_heading(k.replace('_', ' ').title(), level=2)
                doc.add_paragraph(v)
        
        # Generate filename if not provided
        if not filename:
            safe_series = "".join(c for c in input_data.note_series if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"PDS_{safe_series}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Save document
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        print(f"PDS document saved to: {file_path}")
        return file_path

    # ---------------------------------------------------------------------
    # Large Text Templates Rendering (optional flow directly from templates)
    # ---------------------------------------------------------------------
    @staticmethod
    def _extract_unresolved_placeholders(text: str) -> list[str]:
        import re
        pattern = re.compile(r"\[([^\[\]]+)\]")
        return pattern.findall(text or "")

    @classmethod
    def validate_no_unresolved_placeholders(cls, section_contents: Dict[str, str]) -> Dict[str, list[str]]:
        unresolved: Dict[str, list[str]] = {}
        for section_name, content in section_contents.items():
            found = cls._extract_unresolved_placeholders(content)
            if found:
                unresolved[section_name] = sorted(set(found))
        return unresolved

    def create_docx_from_templates(
        self,
        document_sections: Dict[str, str],
        filename: Optional[str] = None,
        title: Optional[str] = None,
        enforce_placeholder_validation: bool = True,
    ) -> str:
        """
        Create a formatted DOCX from canonical PDS large text sections.
        Expected canonical keys follow agents.pds.large_text_templates.list_canonical_section_keys().
        """
        if enforce_placeholder_validation:
            unresolved = self.validate_no_unresolved_placeholders(document_sections)
            if unresolved:
                details = []
                for section_name, missing in unresolved.items():
                    details.append(f"- {section_name}: {', '.join(missing)}")
                    
                raise ValueError(
                    "Unresolved placeholders detected in document sections.\n" + "\n".join(details)
                )

        ordered_sections = [
            "initial_disclaimers",
            "offering_details",
            "nature_and_profile",
            "prospectus_structure",
            "principal_at_risk_notes",
            "calculation_agent_determinations",
            "risk_factor_introduction",
            "potential_for_loss",
        ]
        section_titles = {
            "initial_disclaimers": "Initial Declarations and Disclaimers",
            "offering_details": "Offering Details",
            "nature_and_profile": "Nature of the Notes and Investment Profile",
            "prospectus_structure": "Prospectus Structure and Document Incorporation",
            "principal_at_risk_notes": "Principal at Risk Notes",
            "calculation_agent_determinations": "Determinations of the Calculation Agent",
            "risk_factor_introduction": "Risk Factors Introduction",
            "potential_for_loss": "Potential for Loss",
        }

        if Document is None:
            raise RuntimeError("python-docx is required to create DOCX documents")
        doc = Document()
        if title:
            doc.add_heading(title, level=0)

        for section_key in ordered_sections:
            content = document_sections.get(section_key)
            if not content:
                continue
            doc.add_heading(section_titles[section_key], level=1)
            doc.add_paragraph(content)

        if not filename:
            filename = f"PDS_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        print(f"PDS templates document saved to: {file_path}")
        return file_path