"""
Document generator for BSP (Base Shelf Prospectus) agent outputs.

TODO: Full implementation pending ISM agent completion.
"""

import os
from datetime import datetime
from typing import Optional, Dict, Any, List
from docx import Document

from .models import BSPOutput, BSPInput
from .config import BSPConfig
from core.config import global_config


class BSPDocumentGenerator:
    """
    Generator for creating formatted BSP documents from structured output.
    
    TODO: Implement full document generator following ISM pattern.
    """
    
    def __init__(self, config: Optional[BSPConfig] = None):
        """
        Initialize the document generator.
        
        Args:
            config: BSP configuration for formatting preferences
        """
        self.config = config or BSPConfig.get_default_config()
        self.output_dir = global_config.get_output_path("bsp")
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_docx_document(
        self,
        bsp_output: BSPOutput,
        input_data: BSPInput,
        filename: Optional[str] = None,
    ) -> str:
        """
        Create a formatted DOCX document from BSP structured output.

        Args:
            bsp_output: Structured output from BSP agent
            input_data: Original input data for context
            filename: Custom filename (optional)

        Returns:
            Path to the created document
        """
        document = Document()

        # Title
        document.add_heading(bsp_output.document_title, level=0)

        # Cover Page
        document.add_heading("Cover Page", level=1)
        document.add_paragraph(bsp_output.cover_page)

        # Issuer Information
        document.add_heading("Issuer Information", level=1)
        document.add_paragraph(bsp_output.issuer_description)
        document.add_heading("Business Overview", level=2)
        document.add_paragraph(bsp_output.business_overview)
        document.add_heading("Financial Information", level=2)
        document.add_paragraph(bsp_output.financial_information)

        # Program Description
        document.add_heading("Program Overview", level=1)
        document.add_paragraph(bsp_output.program_overview)
        document.add_heading("General Terms", level=2)
        document.add_paragraph(bsp_output.general_terms)

        # Risk Factors
        document.add_heading("Risk Factors", level=1)
        document.add_paragraph(bsp_output.risk_factors)

        # Legal Terms
        document.add_heading("Legal Terms", level=1)
        document.add_paragraph(bsp_output.legal_terms)

        # Use of Proceeds
        document.add_heading("Use of Proceeds", level=1)
        document.add_paragraph(bsp_output.use_of_proceeds)

        # Regulatory Disclosures
        document.add_heading("Regulatory Disclosures", level=1)
        document.add_paragraph(bsp_output.regulatory_disclosures)

        # Additional Sections
        if bsp_output.additional_sections:
            for section_name, section_text in bsp_output.additional_sections.items():
                document.add_heading(section_name, level=1)
                document.add_paragraph(section_text)

        # Footer metadata
        document.add_paragraph("")
        document.add_paragraph(f"Document Version: {bsp_output.document_version}")
        document.add_paragraph(f"Generation Date: {bsp_output.generation_date}")

        # Filename
        if not filename:
            safe_program_name = "".join(
                c for c in input_data.program_name if c.isalnum() or c in (" ", "-", "_")
            ).rstrip()
            filename = f"BSP_{safe_program_name}_{datetime.now().strftime('%Y%m%d')}.docx"

        # Save
        file_path = os.path.join(self.output_dir, filename)
        document.save(file_path)
        print(f"BSP document saved to: {file_path}")
        return file_path

    # ---------------------------------------------------------------------
    # Large Text Templates Rendering
    # ---------------------------------------------------------------------
    @staticmethod
    def _extract_unresolved_placeholders(text: str) -> List[str]:
        """
        Extract unresolved placeholders of the form [PLACEHOLDER] in a text blob.
        Returns the placeholder names as a list (without brackets).
        """
        import re

        pattern = re.compile(r"\[([^\[\]]+)\]")
        return pattern.findall(text or "")

    @classmethod
    def validate_no_unresolved_placeholders(
        cls, section_contents: Dict[str, str]
    ) -> Dict[str, List[str]]:
        """
        Validate that no unresolved placeholders remain in section contents.

        Returns a dict mapping section_name -> list of unresolved placeholders.
        The caller can treat a non-empty mapping as a validation error.
        """
        unresolved: Dict[str, List[str]] = {}
        for section_name, content in section_contents.items():
            found = cls._extract_unresolved_placeholders(content)
            if found:
                unresolved[section_name] = sorted(set(found))
        return unresolved

    def create_docx_from_templates(
        self,
        document_sections: Dict[str, str],
        filename: Optional[str] = None,
        title: Optional[str] = None,
        enforce_placeholder_validation: bool = True,
    ) -> str:
        """
        Create a formatted DOCX document from canonical BSP large text sections.

        Args:
            document_sections: Mapping of canonical section keys to rendered text
                Keys should be from the canonical set used by large text templates:
                - cover_page_disclosures
                - forward_looking_statements
                - documents_incorporated_by_reference
                - description_of_the_notes
                - plan_of_distribution
                - risk_factors
                - use_of_proceeds
                - purchasers_statutory_rights
                - certificate_of_the_bank
                - certificate_of_the_dealers
            filename: Optional custom filename
            title: Optional main document title (heading level 0)
            enforce_placeholder_validation: If True, will raise ValueError if any
                unresolved placeholders like [PLACEHOLDER] are detected in content

        Returns:
            Path to the created DOCX file
        """
        # Validate placeholders
        if enforce_placeholder_validation:
            unresolved = self.validate_no_unresolved_placeholders(document_sections)
            if unresolved:
                details = []
                for section_name, missing in unresolved.items():
                    details.append(f"- {section_name}: {', '.join(missing)}")
                message = (
                    "Unresolved placeholders detected in document sections.\n" +
                    "\n".join(details)
                )
                raise ValueError(message)

        # Friendly titles and order
        ordered_sections: List[str] = [
            "cover_page_disclosures",
            "forward_looking_statements",
            "documents_incorporated_by_reference",
            "description_of_the_notes",
            "plan_of_distribution",
            "risk_factors",
            "use_of_proceeds",
            "purchasers_statutory_rights",
            "certificate_of_the_bank",
            "certificate_of_the_dealers",
        ]
        section_titles: Dict[str, str] = {
            "cover_page_disclosures": "Cover Page Disclosures",
            "forward_looking_statements": "Forward-Looking Statements",
            "documents_incorporated_by_reference": "Documents Incorporated by Reference",
            "description_of_the_notes": "Description of the Notes",
            "plan_of_distribution": "Plan of Distribution",
            "risk_factors": "Risk Factors",
            "use_of_proceeds": "Use of Proceeds",
            "purchasers_statutory_rights": "Purchaser's Statutory Rights",
            "certificate_of_the_bank": "Certificate of the Bank",
            "certificate_of_the_dealers": "Certificate of the Dealers",
        }

        document = Document()

        # Title
        if title:
            document.add_heading(title, level=0)

        # Sections
        for section_key in ordered_sections:
            content = document_sections.get(section_key)
            if not content:
                # Skip missing sections silently to allow partial documents
                continue
            document.add_heading(section_titles[section_key], level=1)
            document.add_paragraph(content)

        # Filename default
        if not filename:
            filename = f"BSP_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Save
        file_path = os.path.join(self.output_dir, filename)
        document.save(file_path)
        print(f"BSP templates document saved to: {file_path}")
        return file_path

    def save_sections(
        self,
        document_sections: Dict[str, str],
        filename_stem: Optional[str] = None,
        title: Optional[str] = None,
    ) -> Dict[str, str]:
        """
        Save BSP large-text sections to JSON and TXT.

        If filename_stem is not provided, a timestamped default is used. Prefer passing
        the DOCX stem for consistent naming across formats.
        """
        import json
        if not filename_stem:
            filename_stem = f"BSP_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        json_path = os.path.join(self.output_dir, f"{filename_stem}.json")
        txt_path = os.path.join(self.output_dir, f"{filename_stem}.txt")

        # JSON dump of section map
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(document_sections, f, ensure_ascii=False, indent=2)

        # TXT rendering in canonical order
        ordered_sections: List[str] = [
            "cover_page_disclosures",
            "forward_looking_statements",
            "documents_incorporated_by_reference",
            "description_of_the_notes",
            "plan_of_distribution",
            "risk_factors",
            "use_of_proceeds",
            "purchasers_statutory_rights",
            "certificate_of_the_bank",
            "certificate_of_the_dealers",
        ]
        section_titles: Dict[str, str] = {
            "cover_page_disclosures": "Cover Page Disclosures",
            "forward_looking_statements": "Forward-Looking Statements",
            "documents_incorporated_by_reference": "Documents Incorporated by Reference",
            "description_of_the_notes": "Description of the Notes",
            "plan_of_distribution": "Plan of Distribution",
            "risk_factors": "Risk Factors",
            "use_of_proceeds": "Use of Proceeds",
            "purchasers_statutory_rights": "Purchaser's Statutory Rights",
            "certificate_of_the_bank": "Certificate of the Bank",
            "certificate_of_the_dealers": "Certificate of the Dealers",
        }

        with open(txt_path, "w", encoding="utf-8") as f:
            if title:
                f.write(title + "\n\n")
            for key in ordered_sections:
                content = document_sections.get(key)
                if not content:
                    continue
                f.write(f"[{section_titles[key]}]\n")
                f.write(content.strip() + "\n\n")

        return {"json_path": json_path, "txt_path": txt_path}