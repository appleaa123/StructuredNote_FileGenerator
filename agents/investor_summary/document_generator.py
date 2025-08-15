"""
Document generator for ISM (Investor Summary) agent outputs.
"""

import os
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .models import ISMOutput, ISMInput
from .config import ISMConfig
from core.config import global_config


class ISMDocumentGenerator:
    """
    Generator for creating formatted ISM documents from structured output.
    
    This class takes the structured output from the ISM agent and converts
    it into professionally formatted documents in various formats (DOCX, PDF, etc.).
    """
    
    def __init__(self, config: Optional[ISMConfig] = None):
        """
        Initialize the document generator.
        
        Args:
            config: ISM configuration for formatting preferences
        """
        self.config = config or ISMConfig.get_default_config()
        self.output_dir = global_config.get_output_path("ism")
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_docx_document(
        self, 
        ism_output: ISMOutput, 
        input_data: ISMInput,
        filename: Optional[str] = None,
        include_metadata: bool = True
    ) -> str:
        """
        Create a formatted DOCX document from ISM output.
        
        Args:
            ism_output: Structured output from ISM agent
            input_data: Original input data for context
            filename: Custom filename (optional)
            include_metadata: Whether to include document metadata
            
        Returns:
            Path to the created document
        """
        # Create new document
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Add document header
        self._add_document_header(doc, ism_output, input_data)
        
        # Add executive summary
        self._add_executive_summary(doc, ism_output)
        
        # Add product overview section
        self._add_product_overview(doc, ism_output)
        
        # Add investment information section
        self._add_investment_information(doc, ism_output)
        
        # Add risk information section
        self._add_risk_information(doc, ism_output)
        
        # Add important information section
        self._add_important_information(doc, ism_output)
        
        # Add regulatory and legal section
        self._add_regulatory_section(doc, ism_output)
        
        # Add contact and support section
        self._add_contact_section(doc, ism_output)
        
        # Add footer information
        self._add_footer_information(doc, ism_output)
        
        # Generate filename if not provided
        if not filename:
            safe_product_name = "".join(c for c in input_data.product_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_issuer = "".join(c for c in input_data.issuer if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"ISM_{safe_issuer}_{safe_product_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Save document
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        print(f"ISM document saved to: {file_path}")
        return file_path
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document"""
        
        # Title style
        title_style = doc.styles.add_style('ISM Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        heading1_style = doc.styles.add_style('ISM Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.font.name = 'Arial'
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        body_style = doc.styles.add_style('ISM Body', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
        
        # Warning style
        warning_style = doc.styles.add_style('ISM Warning', WD_STYLE_TYPE.PARAGRAPH)
        warning_style.font.name = 'Arial'
        warning_style.font.size = Pt(11)
        warning_style.font.bold = True
        warning_style.paragraph_format.space_after = Pt(6)
    
    def _add_document_header(self, doc: Document, ism_output: ISMOutput, input_data: ISMInput):
        """Add document header with title and key information"""
        
        # Document title
        title_paragraph = doc.add_paragraph(ism_output.document_title, style='ISM Title')
        
        # Add key product information table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Add key information rows
        info_items = [
            ("Issuer", input_data.issuer),
            ("Product Type", input_data.product_type.replace('_', ' ').title()),
            ("Underlying Asset", input_data.underlying_asset),
            ("Currency", input_data.currency),
            ("Issue Date", input_data.issue_date.strftime('%B %d, %Y')),
            ("Maturity Date", input_data.maturity_date.strftime('%B %d, %Y')),
            ("Investment Period", f"{((input_data.maturity_date - input_data.issue_date).days / 365.25):.1f} years")
        ]
        
        for label, value in info_items:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
            row[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
    
    def _add_executive_summary(self, doc: Document, ism_output: ISMOutput):
        """Add executive summary section"""
        
        doc.add_paragraph("Executive Summary", style='ISM Heading 1')
        doc.add_paragraph(ism_output.executive_summary, style='ISM Body')
        doc.add_paragraph()
    
    def _add_product_overview(self, doc: Document, ism_output: ISMOutput):
        """Add product overview section"""
        
        doc.add_paragraph("Product Overview", style='ISM Heading 1')
        
        # Product description
        doc.add_paragraph("What is this investment?", style='Heading 2')
        doc.add_paragraph(ism_output.product_description, style='ISM Body')
        
        # How it works
        doc.add_paragraph("How does it work?", style='Heading 2')
        doc.add_paragraph(ism_output.how_it_works, style='ISM Body')
        
        # Key features
        doc.add_paragraph("Key Features", style='Heading 2')
        if isinstance(ism_output.key_features, list):
            for feature in ism_output.key_features:
                p = doc.add_paragraph(style='ISM Body')
                p.add_run("• ").font.bold = True
                p.add_run(feature)
        else:
            doc.add_paragraph(ism_output.key_features, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_investment_information(self, doc: Document, ism_output: ISMOutput):
        """Add investment information section"""
        
        doc.add_paragraph("Investment Information", style='ISM Heading 1')
        
        # Investment details
        doc.add_paragraph("Investment Details", style='Heading 2')
        doc.add_paragraph(ism_output.investment_details, style='ISM Body')
        
        # Potential returns
        doc.add_paragraph("Potential Returns", style='Heading 2')
        doc.add_paragraph(ism_output.potential_returns, style='ISM Body')
        
        # Scenario analysis
        doc.add_paragraph("Market Scenarios", style='Heading 2')
        doc.add_paragraph(ism_output.scenarios_analysis, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_risk_information(self, doc: Document, ism_output: ISMOutput):
        """Add risk information section"""
        
        doc.add_paragraph("Risk Information", style='ISM Heading 1')
        
        # Risk level indicator (highlighted)
        risk_paragraph = doc.add_paragraph(style='ISM Warning')
        risk_paragraph.add_run("Risk Level: ").font.bold = True
        risk_paragraph.add_run(ism_output.risk_level_indicator)
        
        # Risk summary
        doc.add_paragraph("Risk Summary", style='Heading 2')
        doc.add_paragraph(ism_output.risk_summary, style='ISM Body')
        
        # Key risks
        doc.add_paragraph("Key Risks", style='Heading 2')
        if isinstance(ism_output.key_risks, list):
            for risk in ism_output.key_risks:
                p = doc.add_paragraph(style='ISM Body')
                p.add_run("⚠ ").font.bold = True
                p.add_run(risk)
        else:
            doc.add_paragraph(ism_output.key_risks, style='ISM Body')
        
        # Risk mitigation
        doc.add_paragraph("Risk Management", style='Heading 2')
        doc.add_paragraph(ism_output.risk_mitigation, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_important_information(self, doc: Document, ism_output: ISMOutput):
        """Add important information section"""
        
        doc.add_paragraph("Important Information", style='ISM Heading 1')
        
        # Important dates
        doc.add_paragraph("Important Dates", style='Heading 2')
        doc.add_paragraph(ism_output.important_dates, style='ISM Body')
        
        # Fees and charges
        doc.add_paragraph("Fees and Charges", style='Heading 2')
        doc.add_paragraph(ism_output.fees_and_charges, style='ISM Body')
        
        # Liquidity information
        doc.add_paragraph("Liquidity", style='Heading 2')
        doc.add_paragraph(ism_output.liquidity_information, style='ISM Body')
        
        # Suitability
        doc.add_paragraph("Suitability", style='Heading 2')
        doc.add_paragraph(ism_output.suitability_assessment, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_regulatory_section(self, doc: Document, ism_output: ISMOutput):
        """Add regulatory and legal information"""
        
        doc.add_paragraph("Regulatory Information", style='ISM Heading 1')
        
        # Regulatory notices
        doc.add_paragraph("Regulatory Notices", style='Heading 2')
        doc.add_paragraph(ism_output.regulatory_notices, style='ISM Body')
        
        # Tax considerations
        doc.add_paragraph("Tax Considerations", style='Heading 2')
        doc.add_paragraph(ism_output.tax_considerations, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_contact_section(self, doc: Document, ism_output: ISMOutput):
        """Add contact and support information"""
        
        doc.add_paragraph("Contact & Support", style='ISM Heading 1')
        
        # Contact information
        doc.add_paragraph("Contact Information", style='Heading 2')
        doc.add_paragraph(ism_output.contact_information, style='ISM Body')
        
        # Next steps
        doc.add_paragraph("Next Steps", style='Heading 2')
        doc.add_paragraph(ism_output.next_steps, style='ISM Body')
        
        doc.add_paragraph()
    
    def _add_footer_information(self, doc: Document, ism_output: ISMOutput):
        """Add footer information including disclaimers"""
        
        # Add page break before disclaimers
        doc.add_page_break()
        
        doc.add_paragraph("Important Disclaimers", style='ISM Heading 1')
        doc.add_paragraph(ism_output.disclaimer, style='ISM Body')
        
        # Document metadata
        doc.add_paragraph()
        metadata_paragraph = doc.add_paragraph(style='ISM Body')
        metadata_paragraph.add_run(f"Document Version: {ism_output.document_version} | ")
        metadata_paragraph.add_run(f"Generated: {ism_output.generation_date}")
        metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_summary_report(
        self, 
        ism_output: ISMOutput, 
        input_data: ISMInput,
        filename: Optional[str] = None
    ) -> str:
        """
        Create a shorter summary report version.
        
        Args:
            ism_output: Structured output from ISM agent
            input_data: Original input data
            filename: Custom filename (optional)
            
        Returns:
            Path to the created summary report
        """
        doc = Document()
        
        # Set up basic styles
        self._setup_document_styles(doc)
        
        # Title
        doc.add_paragraph(f"Investment Summary: {input_data.product_name}", style='ISM Title')
        
        # Key information table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        key_info = [
            ("Issuer", input_data.issuer),
            ("Underlying", input_data.underlying_asset),
            ("Term", f"{((input_data.maturity_date - input_data.issue_date).days / 365.25):.1f} years"),
            ("Risk Level", ism_output.risk_level_indicator)
        ]
        
        for label, value in key_info:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)
            row[0].paragraphs[0].runs[0].font.bold = True
        
        # Executive summary
        doc.add_paragraph()
        doc.add_paragraph("Executive Summary", style='ISM Heading 1')
        doc.add_paragraph(ism_output.executive_summary, style='ISM Body')
        
        # Key risks (abbreviated)
        doc.add_paragraph("Key Risks", style='ISM Heading 1')
        doc.add_paragraph(ism_output.risk_summary, style='ISM Body')
        
        # Generate filename
        if not filename:
            safe_name = "".join(c for c in input_data.product_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"ISM_Summary_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Save document
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        print(f"ISM summary report saved to: {file_path}")
        return file_path
    
    # ---------------------------------------------------------------------
    # Large Text Templates Rendering (optional flow directly from templates)
    # ---------------------------------------------------------------------
    @staticmethod
    def _extract_unresolved_placeholders(text: str) -> list[str]:
        import re
        pattern = re.compile(r"\[([^\[\]]+)\]")
        return pattern.findall(text or "")
    
    @classmethod
    def validate_no_unresolved_placeholders(cls, section_contents: Dict[str, str]) -> Dict[str, list[str]]:
        unresolved: Dict[str, list[str]] = {}
        for section_name, content in section_contents.items():
            found = cls._extract_unresolved_placeholders(content)
            if found:
                unresolved[section_name] = sorted(set(found))
        return unresolved
    
    def create_docx_from_templates(
        self,
        document_sections: Dict[str, str],
        filename: Optional[str] = None,
        title: Optional[str] = None,
        enforce_placeholder_validation: bool = True,
    ) -> str:
        """
        Create a formatted DOCX from ISM large text template sections.
        Expected keys include: 'executive_summary', 'key_terms', 'additional_key_terms',
        'scenarios', and 'disclaimer'. Extra keys are ignored.
        """
        if enforce_placeholder_validation:
            unresolved = self.validate_no_unresolved_placeholders(document_sections)
            if unresolved:
                details = []
                for section_name, missing in unresolved.items():
                    details.append(f"- {section_name}: {', '.join(missing)}")
                raise ValueError(
                    "Unresolved placeholders detected in document sections.\n" + "\n".join(details)
                )
        
        # Create document
        doc = Document()
        self._setup_document_styles(doc)
        
        if title:
            doc.add_paragraph(title, style='ISM Title')
        
        # Define rendering order
        ordered_sections = [
            ("Executive Summary", "executive_summary"),
            ("Key Terms", "key_terms"),
            ("Additional Key Terms", "additional_key_terms"),
            ("Scenarios", "scenarios"),
            ("Disclaimer", "disclaimer"),
        ]
        
        for heading, key in ordered_sections:
            content = document_sections.get(key)
            if not content:
                continue
            doc.add_paragraph(heading, style='ISM Heading 1')
            doc.add_paragraph(content, style='ISM Body')
            doc.add_paragraph()
        
        if not filename:
            filename = f"ISM_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        print(f"ISM templates document saved to: {file_path}")
        return file_path

    def save_sections(
        self,
        document_sections: Dict[str, str],
        filename_stem: Optional[str] = None,
        title: Optional[str] = None,
    ) -> Dict[str, str]:
        """
        Save ISM large-text sections to JSON and TXT alongside DOCX flows.

        If no filename_stem is provided, a timestamped default is used. The caller
        should normally pass the DOCX filename stem for consistent naming.
        """
        import json

        if not filename_stem:
            filename_stem = f"ISM_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        json_path = os.path.join(self.output_dir, f"{filename_stem}.json")
        txt_path = os.path.join(self.output_dir, f"{filename_stem}.txt")

        # JSON
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(document_sections, f, ensure_ascii=False, indent=2)

        # TXT (ordered, human-readable)
        ordered_sections = [
            ("Executive Summary", "executive_summary"),
            ("Key Terms", "key_terms"),
            ("Additional Key Terms", "additional_key_terms"),
            ("Scenarios", "scenarios"),
            ("Disclaimer", "disclaimer"),
        ]

        with open(txt_path, "w", encoding="utf-8") as f:
            if title:
                f.write(title + "\n\n")
            for heading, key in ordered_sections:
                content = document_sections.get(key)
                if not content:
                    continue
                f.write(f"[{heading}]\n")
                f.write(content.strip() + "\n\n")

        return {"json_path": json_path, "txt_path": txt_path}
    
    def get_document_metadata(self, ism_output: ISMOutput, input_data: ISMInput) -> Dict[str, Any]:
        """
        Get metadata about the generated document.
        
        Args:
            ism_output: Generated ISM output
            input_data: Original input data
            
        Returns:
            Dictionary containing document metadata
        """
        return {
            "document_type": "Investor Summary (ISM)",
            "product_name": input_data.product_name,
            "issuer": input_data.issuer,
            "underlying_asset": input_data.underlying_asset,
            "target_audience": input_data.target_audience,
            "risk_level": ism_output.risk_level_indicator,
            "document_version": ism_output.document_version,
            "generation_date": ism_output.generation_date,
            "file_format": "DOCX",
            "agent_type": "ISM",
            "config_used": self.config.model_dump()
        }