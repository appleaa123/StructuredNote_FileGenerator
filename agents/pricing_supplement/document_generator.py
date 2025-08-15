"""
Document generator for PRS (Pricing Supplement) agent outputs.

TODO: Full implementation pending ISM agent completion.
"""

import os
from datetime import datetime
from typing import Optional, Dict, Any
try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional dependency in CI
    Document = None  # type: ignore

from .models import PRSOutput, PRSInput
from .config import PRSConfig
from core.config import global_config


class PRSDocumentGenerator:
    """
    Generator for creating formatted PRS documents from structured output.
    
    TODO: Implement full document generator following ISM pattern.
    """
    
    def __init__(self, config: Optional[PRSConfig] = None):
        """
        Initialize the document generator.
        
        Args:
            config: PRS configuration for formatting preferences
        """
        self.config = config or PRSConfig.get_default_config()
        self.output_dir = global_config.get_output_path("prs")
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_docx_document(
        self, 
        prs_output: PRSOutput, 
        input_data: PRSInput,
        filename: Optional[str] = None
    ) -> str:
        """
        Create a formatted DOCX document from PRS output.
        
        TODO: Implement comprehensive document creation following ISM pattern.
        
        Args:
            prs_output: Structured output from PRS agent
            input_data: Original input data for context
            filename: Custom filename (optional)
            
        Returns:
            Path to the created document
        """
        if Document is None:
            raise RuntimeError("python-docx is required to create DOCX documents")
        doc = Document()

        # Title
        doc.add_heading(prs_output.document_title, level=0)

        # Pricing Summary
        doc.add_heading("Pricing Summary", level=1)
        doc.add_paragraph(prs_output.pricing_summary)

        # Final Terms Table
        doc.add_heading("Final Terms", level=1)
        doc.add_paragraph(prs_output.final_terms_table)

        # References
        doc.add_heading("Document References", level=1)
        doc.add_paragraph(prs_output.document_references)

        # Pricing Methodology and Estimated Value
        doc.add_heading("Pricing Methodology", level=1)
        doc.add_paragraph(prs_output.pricing_methodology)
        doc.add_heading("Estimated Value", level=2)
        doc.add_paragraph(prs_output.estimated_value_explanation)

        # Settlement and Delivery
        doc.add_heading("Settlement Instructions", level=1)
        doc.add_paragraph(prs_output.settlement_instructions)
        doc.add_heading("Delivery Procedures", level=2)
        doc.add_paragraph(prs_output.delivery_procedures)

        # Distribution and Fees
        doc.add_heading("Distribution Information", level=1)
        doc.add_paragraph(prs_output.distribution_information)
        doc.add_heading("Fees and Expenses", level=2)
        doc.add_paragraph(prs_output.fees_and_expenses)

        # Market Data
        doc.add_heading("Market Data at Pricing", level=1)
        doc.add_paragraph(prs_output.market_data_at_pricing)

        # Regulatory Notices
        doc.add_heading("Regulatory Notices", level=1)
        doc.add_paragraph(prs_output.regulatory_notices)

        # Additional Sections
        if prs_output.additional_sections:
            for section_name, section_text in prs_output.additional_sections.items():
                doc.add_heading(section_name.replace('_', ' ').title(), level=1)
                doc.add_paragraph(section_text)

        # Footer metadata
        doc.add_paragraph("")
        doc.add_paragraph(f"Document Version: {prs_output.document_version}")
        doc.add_paragraph(f"Generation Date: {prs_output.generation_date}")
        doc.add_paragraph(f"Pricing Timestamp: {prs_output.pricing_timestamp}")
        
        # Generate filename if not provided
        if not filename:
            pricing_date_str = input_data.pricing_date.strftime('%Y%m%d')
            filename = f"PRS_{pricing_date_str}_{datetime.now().strftime('%H%M')}.docx"
        
        # Save document
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        
        print(f"PRS document saved to: {file_path}")
        return file_path

    # ---------------------------------------------------------------------
    # Large Text Templates Rendering (optional flow directly from templates)
    # ---------------------------------------------------------------------
    @staticmethod
    def _extract_unresolved_placeholders(text: str) -> list[str]:
        import re
        pattern = re.compile(r"\[([^\[\]]+)\]")
        return pattern.findall(text or "")

    @classmethod
    def validate_no_unresolved_placeholders(cls, section_contents: Dict[str, str]) -> Dict[str, list[str]]:
        unresolved: Dict[str, list[str]] = {}
        for section_name, content in section_contents.items():
            found = cls._extract_unresolved_placeholders(content)
            if found:
                unresolved[section_name] = sorted(set(found))
        return unresolved

    def create_docx_from_templates(
        self,
        document_sections: Dict[str, str],
        filename: Optional[str] = None,
        title: Optional[str] = None,
        enforce_placeholder_validation: bool = True,
    ) -> str:
        """
        Create a formatted DOCX from canonical PRS large text sections.
        Expected canonical keys follow agents.prs.large_text_templates.list_canonical_section_keys().
        """
        if enforce_placeholder_validation:
            unresolved = self.validate_no_unresolved_placeholders(document_sections)
            if unresolved:
                details = []
                for section_name, missing in unresolved.items():
                    details.append(f"- {section_name}: {', '.join(missing)}")
                raise ValueError(
                    "Unresolved placeholders detected in document sections.\n" + "\n".join(details)
                )

        ordered_sections = [
            "regulatory_and_offering_disclaimers",
            "offering_overview",
            "general_risks_and_guarantees",
            "prospectus_and_capitalized_terms",
            "documents_incorporated_by_reference",
            "deferred_payment",
            "forward_looking_statements",
            "suitability_for_investment",
            "appendix_c_certain_canadian_federal_income_tax_considerations",
        ]
        section_titles = {
            "regulatory_and_offering_disclaimers": "Regulatory and Offering Disclaimers",
            "offering_overview": "Offering Overview",
            "general_risks_and_guarantees": "General Risks and Guarantees",
            "prospectus_and_capitalized_terms": "Prospectus and Capitalized Terms",
            "documents_incorporated_by_reference": "Documents Incorporated by Reference",
            "deferred_payment": "Deferred Payment",
            "forward_looking_statements": "Forward-Looking Statements",
            "suitability_for_investment": "Suitability for Investment",
            "appendix_c_certain_canadian_federal_income_tax_considerations": "Appendix C: Certain Canadian Federal Income Tax Considerations",
        }

        if Document is None:
            raise RuntimeError("python-docx is required to create DOCX documents")
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        for section_key in ordered_sections:
            content = document_sections.get(section_key)
            if not content:
                continue
            doc.add_heading(section_titles[section_key], level=1)
            doc.add_paragraph(content)

        if not filename:
            filename = f"PRS_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        print(f"PRS templates document saved to: {file_path}")
        return file_path

    def save_sections(
        self,
        document_sections: Dict[str, str],
        filename_stem: Optional[str] = None,
        title: Optional[str] = None,
    ) -> Dict[str, str]:
        """
        Save PRS large-text sections to JSON and TXT.

        If filename_stem is not provided, a timestamp-based default is used. Prefer passing
        the DOCX stem for consistent naming across formats.
        """
        import json
        if not filename_stem:
            filename_stem = f"PRS_Templates_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        json_path = os.path.join(self.output_dir, f"{filename_stem}.json")
        txt_path = os.path.join(self.output_dir, f"{filename_stem}.txt")

        # JSON dump of section map
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(document_sections, f, ensure_ascii=False, indent=2)

        ordered_sections = [
            "regulatory_and_offering_disclaimers",
            "offering_overview",
            "general_risks_and_guarantees",
            "prospectus_and_capitalized_terms",
            "documents_incorporated_by_reference",
            "deferred_payment",
            "forward_looking_statements",
            "suitability_for_investment",
            "appendix_c_certain_canadian_federal_income_tax_considerations",
        ]
        section_titles = {
            "regulatory_and_offering_disclaimers": "Regulatory and Offering Disclaimers",
            "offering_overview": "Offering Overview",
            "general_risks_and_guarantees": "General Risks and Guarantees",
            "prospectus_and_capitalized_terms": "Prospectus and Capitalized Terms",
            "documents_incorporated_by_reference": "Documents Incorporated by Reference",
            "deferred_payment": "Deferred Payment",
            "forward_looking_statements": "Forward-Looking Statements",
            "suitability_for_investment": "Suitability for Investment",
            "appendix_c_certain_canadian_federal_income_tax_considerations": "Appendix C: Certain Canadian Federal Income Tax Considerations",
        }

        # TXT rendering
        with open(txt_path, "w", encoding="utf-8") as f:
            if title:
                f.write(title + "\n\n")
            for key in ordered_sections:
                content = document_sections.get(key)
                if not content:
                    continue
                f.write(f"[{section_titles[key]}]\n")
                f.write(content.strip() + "\n\n")

        return {"json_path": json_path, "txt_path": txt_path}